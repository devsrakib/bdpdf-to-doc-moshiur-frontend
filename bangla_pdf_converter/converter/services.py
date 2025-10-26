import os
import re
from django.conf import settings
# import google.generativeai as genai
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx.shared import Pt, Inches
from django.core.files.base import ContentFile

# Configure Tesseract if custom path is set
if hasattr(settings, 'TESSERACT_CMD'):
    pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD

# # Configure Gemini (COMMENTED OUT FOR NOW)
# genai.configure(api_key=settings.GEMINI_API_KEY)


class PDFProcessingService:
    """Service class for PDF processing operations"""
    
    def __init__(self):
        # self.model = genai.GenerativeModel('gemini-1.5-flash')
        pass
    
    def extract_text_from_pdf(self, pdf_path):
        """Extract text from Bangla PDF using OCR"""
        print(f"Converting PDF to images: {pdf_path}")
        
        try:
            # Convert PDF pages to images
            images = convert_from_path(pdf_path, dpi=300)
            
            extracted_text = []
            total_pages = len(images)
            
            print(f"Processing {total_pages} pages...")
            
            for i, image in enumerate(images):
                print(f"Processing page {i+1}/{total_pages}...")
                
                # Perform OCR with Bangla language support
                # Use 'ben' for Bengali/Bangla language
                text = pytesseract.image_to_string(image, lang='ben+eng')
                
                extracted_text.append({
                    'page': i + 1,
                    'text': text.strip()
                })
            
            return extracted_text, total_pages
        
        except Exception as e:
            print(f"OCR Error: {str(e)}")
            import traceback
            traceback.print_exc()
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    def proofread_with_ai(self, text):
        """Use Gemini AI to proofread and improve Bangla text (COMMENTED OUT)"""
        # GEMINI PROOFREADING - COMMENTED OUT FOR NOW
        # try:
        #     prompt = f"""You are a Bangla language expert. Please proofread and improve the following Bangla text extracted from OCR.
        #
        # Fix any:
        # - OCR recognition errors
        # - Spelling mistakes
        # - Grammar issues
        # - Formatting inconsistencies
        #
        # Preserve the original structure and formatting as much as possible. Return only the corrected text without any explanations.
        #
        # Text to proofread:
        # {text}
        # """
        #     
        #     response = self.model.generate_content(prompt)
        #     return response.text
        # 
        # except Exception as e:
        #     print(f"AI proofreading error: {e}")
        #     return text
        
        # For now, return the original text without AI proofreading
        return text
    
    def detect_formatting(self, text):
        """Detect titles, subtitles, and paragraph structure"""
        lines = text.split('\n')
        formatted_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                formatted_lines.append({'type': 'empty', 'text': ''})
                continue
            
            # Detect titles (short lines, often centered)
            if len(line) < 50 and len(line.split()) <= 8:
                formatted_lines.append({'type': 'title', 'text': line})
            # Detect subtitles (medium length, numbered)
            elif re.match(r'^[\d]+\.', line) or len(line) < 100:
                formatted_lines.append({'type': 'subtitle', 'text': line})
            else:
                formatted_lines.append({'type': 'paragraph', 'text': line})
        
        return formatted_lines
    
    def create_docx(self, text_data, output_path):
        """Create a formatted Word document"""
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        for page_data in text_data:
            page_num = page_data['page']
            text = page_data['text']
            
            # Detect formatting
            formatted_lines = self.detect_formatting(text)
            
            for line_data in formatted_lines:
                if line_data['type'] == 'empty':
                    doc.add_paragraph()
                elif line_data['type'] == 'title':
                    p = doc.add_heading(line_data['text'], level=1)
                elif line_data['type'] == 'subtitle':
                    p = doc.add_heading(line_data['text'], level=2)
                else:
                    p = doc.add_paragraph(line_data['text'])
                    p.paragraph_format.line_spacing = 1.5
                    p.paragraph_format.space_after = Pt(6)
            
            # Add page break after each page except the last
            if page_num < len(text_data):
                doc.add_page_break()
        
        doc.save(output_path)
        return output_path
    
    def create_txt(self, text_data, output_path):
        """Create a plain text file"""
        with open(output_path, 'w', encoding='utf-8') as f:
            for page_data in text_data:
                f.write(f"--- Page {page_data['page']} ---\n\n")
                f.write(page_data['text'])
                f.write("\n\n")
        
        return output_path
    
    def process_pdf(self, conversion_obj):
        """Main processing method"""
        try:
            # Update status
            conversion_obj.status = 'processing'
            conversion_obj.save()
            
            # Get PDF path
            pdf_path = conversion_obj.uploaded_file.path
            
            # Extract text
            extracted_text, total_pages = self.extract_text_from_pdf(pdf_path)
            conversion_obj.total_pages = total_pages
            conversion_obj.save()
            
            # Process text (currently without AI proofreading)
            processed_text = []
            word_count = 0
            
            for page_data in extracted_text:
                print(f"Processing page {page_data['page']}...")
                # AI proofreading is disabled for now
                corrected_text = page_data['text']
                
                processed_text.append({
                    'page': page_data['page'],
                    'text': corrected_text
                })
                word_count += len(corrected_text.split())
            
            conversion_obj.word_count = word_count
            conversion_obj.save()
            
            # Create output files
            base_filename = os.path.splitext(conversion_obj.original_filename)[0]
            
            # DOCX
            docx_filename = f"{conversion_obj.id}_{base_filename}.docx"
            docx_path = os.path.join(settings.OUTPUT_FOLDER, docx_filename)
            self.create_docx(processed_text, docx_path)
            
            with open(docx_path, 'rb') as f:
                conversion_obj.docx_file.save(docx_filename, ContentFile(f.read()), save=False)
            
            # TXT
            txt_filename = f"{conversion_obj.id}_{base_filename}.txt"
            txt_path = os.path.join(settings.OUTPUT_FOLDER, txt_filename)
            self.create_txt(processed_text, txt_path)
            
            with open(txt_path, 'rb') as f:
                conversion_obj.txt_file.save(txt_filename, ContentFile(f.read()), save=False)
            
            # Update status
            conversion_obj.status = 'completed'
            conversion_obj.save()
            
            # Clean up temporary files
            if os.path.exists(docx_path):
                os.remove(docx_path)
            if os.path.exists(txt_path):
                os.remove(txt_path)
            
            print(f"Conversion completed successfully: {conversion_obj.id}")
            return conversion_obj
        
        except Exception as e:
            print(f"Processing error: {str(e)}")
            conversion_obj.status = 'failed'
            conversion_obj.error_message = str(e)
            conversion_obj.save()
            raise e